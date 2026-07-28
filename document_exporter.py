"""
document_exporter.py

생성된 제안 초안을 파일로 저장하는 모듈입니다.
Markdown(.md)과 DOCX(.docx) 두 가지 형식을 지원합니다.
"""

from pathlib import Path
from datetime import datetime

import docx

OUTPUT_DIR = Path(__file__).parent / "outputs"

# 초안 딕셔너리의 각 키를 화면/문서에 표시할 한글 제목으로 매핑합니다.
DRAFT_LABELS = {
    "necessity": "참여 필요성",
    "center_role": "우리 센터의 담당 역할",
    "work_details": "세부 수행내용",
    "yearly_plan": "연차별 수행계획",
    "deliverables": "예상 산출물",
    "kpi_draft": "KPI 초안",
    "consortium_role": "컨소시엄 내 역할",
    "expected_effects": "기대효과",
    "open_questions": "추가 확인이 필요한 사항",
}


def build_markdown(rfp_result: dict, selected_role: dict, draft: dict) -> str:
    """RFP 분석 결과 + 선택한 역할 + 제안 초안을 하나의 Markdown 문자열로 합칩니다."""
    lines = []
    lines.append(f"# 제안 초안 - {rfp_result.get('project_name', '확인 필요')}")
    lines.append("")
    lines.append(f"_생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}_")
    lines.append("")

    lines.append("## 1. RFP 핵심 정보")
    lines.append(f"- 사업 목적: {rfp_result.get('purpose', '확인 필요')}")
    lines.append(f"- 발주기관: {rfp_result.get('organization', '확인 필요')}")
    lines.append(f"- 사업기간: {rfp_result.get('duration', '확인 필요')}")
    lines.append(f"- 예산: {rfp_result.get('budget', '확인 필요')}")
    lines.append("")

    lines.append("## 2. 선택한 역할 후보")
    lines.append(f"- 역할명: {selected_role.get('role', '확인 필요')}")
    lines.append(f"- 추천 이유: {selected_role.get('reason', '확인 필요')}")
    lines.append("")

    lines.append("## 3. 우리 센터 담당 제안 초안")
    for key, label in DRAFT_LABELS.items():
        lines.append(f"### {label}")
        lines.append(str(draft.get(key, "확인 필요")))
        lines.append("")

    lines.append("---")
    lines.append("※ 본 문서는 AI가 생성한 초안입니다. `[확인 필요]` 표시 항목은 반드시 담당자가 검토·보완해야 합니다.")

    return "\n".join(lines)


def export_markdown(rfp_result: dict, selected_role: dict, draft: dict, filename: str = None) -> str:
    """Markdown 파일로 저장하고 저장된 경로(문자열)를 반환합니다."""
    OUTPUT_DIR.mkdir(exist_ok=True)
    filename = filename or f"proposal_draft_{_timestamp()}.md"
    file_path = OUTPUT_DIR / filename

    content = build_markdown(rfp_result, selected_role, draft)
    file_path.write_text(content, encoding="utf-8")
    return str(file_path)


def export_docx(rfp_result: dict, selected_role: dict, draft: dict, filename: str = None) -> str:
    """DOCX 파일로 저장하고 저장된 경로(문자열)를 반환합니다."""
    OUTPUT_DIR.mkdir(exist_ok=True)
    filename = filename or f"proposal_draft_{_timestamp()}.docx"
    file_path = OUTPUT_DIR / filename

    document = docx.Document()

    document.add_heading(f"제안 초안 - {rfp_result.get('project_name', '확인 필요')}", level=0)
    document.add_paragraph(f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    document.add_heading("1. RFP 핵심 정보", level=1)
    document.add_paragraph(f"사업 목적: {rfp_result.get('purpose', '확인 필요')}")
    document.add_paragraph(f"발주기관: {rfp_result.get('organization', '확인 필요')}")
    document.add_paragraph(f"사업기간: {rfp_result.get('duration', '확인 필요')}")
    document.add_paragraph(f"예산: {rfp_result.get('budget', '확인 필요')}")

    document.add_heading("2. 선택한 역할 후보", level=1)
    document.add_paragraph(f"역할명: {selected_role.get('role', '확인 필요')}")
    document.add_paragraph(f"추천 이유: {selected_role.get('reason', '확인 필요')}")

    document.add_heading("3. 우리 센터 담당 제안 초안", level=1)
    for key, label in DRAFT_LABELS.items():
        document.add_heading(label, level=2)
        document.add_paragraph(str(draft.get(key, "확인 필요")))

    document.add_paragraph(
        "※ 본 문서는 AI가 생성한 초안입니다. [확인 필요] 표시 항목은 반드시 담당자가 검토·보완해야 합니다."
    )

    document.save(str(file_path))
    return str(file_path)


def _timestamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")
