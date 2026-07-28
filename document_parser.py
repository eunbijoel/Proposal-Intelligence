"""
document_parser.py

업로드된 문서(PDF, DOCX)에서 텍스트를 추출하는 모듈입니다.

이 파일이 하는 일은 딱 한 가지입니다.
- 파일 경로를 받아서 -> "텍스트 조각(청크) 목록"으로 돌려줍니다.

각 텍스트 조각에는 다음 정보가 함께 들어 있습니다.
- 원본 파일명
- 위치 정보 (예: "3페이지" 또는 "문단 12")
- 실제 텍스트 내용

이렇게 위치 정보를 같이 저장해 두면, 나중에 RFP 분석 결과에
"어느 파일, 어느 위치에서 나온 내용인지"를 표시할 수 있습니다.
"""

from pathlib import Path
from pypdf import PdfReader
import docx


def extract_text(file_path: str) -> list[dict]:
    """
    파일 하나를 받아서 텍스트 청크 리스트를 반환합니다.

    반환 형태 예시:
    [
        {"file": "공고문.pdf", "location": "1페이지", "text": "..."},
        {"file": "공고문.pdf", "location": "2페이지", "text": "..."},
    ]

    지원하지 않는 확장자는 빈 리스트를 반환하고,
    호출하는 쪽(app.py)에서 사용자에게 안내 메시지를 보여줍니다.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(path)
    elif ext == ".docx":
        return _extract_docx(path)
    else:
        # 지원하지 않는 형식입니다. 이번 버전은 PDF/DOCX만 지원합니다.
        return []


def _extract_pdf(path: Path) -> list[dict]:
    """PDF 파일을 페이지 단위로 텍스트를 추출합니다."""
    chunks = []
    try:
        reader = PdfReader(str(path))
    except Exception as e:
        # 손상된 PDF이거나 암호화된 경우 여기로 옵니다.
        return [{"file": path.name, "location": "파일 열기 실패", "text": f"[오류] {e}"}]

    for page_num, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            chunks.append({
                "file": path.name,
                "location": f"{page_num}페이지",
                "text": text,
            })

    if not chunks:
        # 텍스트가 하나도 없으면 스캔본(이미지) PDF일 가능성이 큽니다.
        chunks.append({
            "file": path.name,
            "location": "전체",
            "text": "[텍스트 추출 실패 - 스캔된 이미지 PDF일 수 있습니다. OCR은 이번 버전에서 지원하지 않습니다]",
        })
    return chunks


def _extract_docx(path: Path) -> list[dict]:
    """DOCX 파일을 문단(paragraph) 단위로 텍스트를 추출합니다."""
    chunks = []
    try:
        document = docx.Document(str(path))
    except Exception as e:
        return [{"file": path.name, "location": "파일 열기 실패", "text": f"[오류] {e}"}]

    para_num = 0
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            para_num += 1
            chunks.append({
                "file": path.name,
                "location": f"문단 {para_num}",
                "text": text,
            })

    # 표 안의 텍스트도 함께 추출합니다.
    for table_idx, table in enumerate(document.tables, start=1):
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                chunks.append({
                    "file": path.name,
                    "location": f"표 {table_idx}",
                    "text": row_text,
                })

    if not chunks:
        chunks.append({
            "file": path.name,
            "location": "전체",
            "text": "[텍스트를 찾지 못했습니다 - 빈 문서이거나 이미지로만 구성된 문서일 수 있습니다]",
        })
    return chunks


def chunks_to_text(chunks: list[dict], max_chars: int = 12000) -> str:
    """
    LLM 프롬프트에 넣기 좋은 형태로 청크 리스트를 하나의 문자열로 합칩니다.
    각 청크 앞에 [파일명 / 위치] 표시를 붙여서, LLM이 출처를 함께 답하도록 유도합니다.

    max_chars: 프롬프트가 너무 길어지지 않도록 자르는 최대 글자 수입니다.
    """
    parts = []
    total_len = 0
    for chunk in chunks:
        header = f"[출처: {chunk['file']} / {chunk['location']}]"
        piece = f"{header}\n{chunk['text']}\n"
        if total_len + len(piece) > max_chars:
            parts.append("\n...(내용이 길어 이후 부분은 생략되었습니다)...")
            break
        parts.append(piece)
        total_len += len(piece)
    return "\n".join(parts)
